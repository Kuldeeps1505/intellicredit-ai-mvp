"""
Agent 7 — Credit Appraisal Memo (CAM) Generation Agent
Day 4 deliverable.

8 CAM sections:
  1. Executive Summary
  2. Company Background & Business Overview (LLM narrative)
  3. Financial Analysis (LLM narrative, no numbers invented)
  4. GST Reconciliation & ITC Analysis (NEW)
  5. Buyer Concentration Risk (NEW)
  6. Risk Assessment — Five-Cs
  7. Credit Decision & Loan Terms
  8. Counterfactual Recommendations (NEW)
  Appendix A: Chain of Evidence

Exports: PDF (WeasyPrint) + DOCX (python-docx)
Endpoints: GET /api/applications/{id}/cam
           GET /api/applications/{id}/cam/download
"""
from __future__ import annotations
import os, uuid, time, re, json
from datetime import datetime
from pathlib import Path

import anthropic

from app.services.redis_service import get_session, set_session, publish_event
from app.services.db_helpers import log_agent, _AgentSession
from app.models import CAMReport, Application, Company, Financial, FieldProvenance, RiskFlag
from sqlalchemy import select
from app.config import settings
from app.services.llm_service import generate_text

AGENT = "cam_generation"
CAM_OUTPUT_DIR = Path("/tmp/cam_reports")
CAM_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


# ── LLM narrative generator ───────────────────────────────
def generate_section_narrative(section_name: str, context: dict) -> str:
    """Generate narrative for LLM-written sections using Gemini."""
    if not settings.gemini_api_key:
        return f"[{section_name} narrative — set GEMINI_API_KEY for AI-generated content]"

    prompts = {
        "company_background": (
            f"Write a 2-paragraph professional credit analyst summary of this company's "
            f"business background for a Credit Appraisal Memo. "
            f"Company: {context.get('company_name')}, "
            f"Sector: {context.get('sector')}, "
            f"CIN: {context.get('cin')}, GSTIN: {context.get('gstin')}. "
            f"Loan purpose: {context.get('purpose')}. "
            "Write only about the business — do not invent financial figures. "
            "Professional tone. No preamble."
        ),
        "financial_narrative": (
            f"Write a 2-paragraph financial analysis narrative for a Credit Appraisal Memo. "
            f"Key ratios — DSCR: {context.get('dscr', 'N/A')}, "
            f"D/E: {context.get('de_ratio', 'N/A')}, "
            f"Net Profit Margin: {context.get('npm', 'N/A')}, "
            f"EBITDA Margin: {context.get('ebitda_margin', 'N/A')}, "
            f"Current Ratio: {context.get('current_ratio', 'N/A')}. "
            f"Anomaly flags: {context.get('flags', 'None')}. "
            "Analyse the financial health trend. Do not invent figures not provided. "
            "Professional credit analyst tone. No preamble."
        ),
        "risk_narrative": (
            f"Write a 1-paragraph overall credit risk summary for a Credit Appraisal Memo. "
            f"Five-Cs final score: {context.get('final_score')}/100, "
            f"risk category: {context.get('risk_category')}, "
            f"decision: {context.get('decision')}. "
            f"Key risk drivers: {context.get('top_drivers', 'N/A')}. "
            "Summarise the credit risk assessment. Do not invent figures. "
            "Professional tone. No preamble."
        ),
    }

    result = generate_text(prompts.get(section_name, "Write a brief summary."), max_tokens=400)
    return result or f"[{section_name} narrative unavailable]"


# ── HTML CAM template ─────────────────────────────────────
def build_cam_html(data: dict) -> str:
    """Build full CAM HTML — all 8 sections + appendix."""
    company = data.get("company", {})
    risk_scores = data.get("risk_scores", {})
    decision_data = data.get("decision", {})
    gst_recon = data.get("gst_reconciliation", {})
    buyer_conc = data.get("buyer_concentration", {})
    counterfactuals = data.get("counterfactuals", {})
    provenance = data.get("provenance", [])
    flags = data.get("flags", [])
    loan_terms = decision_data.get("loan_terms", {})
    checklist = decision_data.get("rbi_compliance_checklist", [])
    ratios = data.get("latest_ratios", {})
    narratives = data.get("narratives", {})

    rec = risk_scores.get("decision", "UNKNOWN")
    rec_color = {"APPROVE": "#16a34a", "CONDITIONAL_APPROVAL": "#d97706",
                 "CONDITIONAL": "#d97706", "REJECT": "#dc2626"}.get(rec, "#6b7280")

    def fmt_flag_row(f):
        col = {"CRITICAL": "#dc2626", "HIGH": "#d97706",
               "MEDIUM": "#2563eb", "LOW": "#6b7280"}.get(f.get("severity"), "#6b7280")
        return (f'<tr><td style="color:{col};font-weight:bold">{f.get("severity")}</td>'
                f'<td>{f.get("flag_type")}</td><td>{f.get("description", "")}</td></tr>')

    def fmt_prov_row(p):
        return (f'<tr><td>{p.get("field_name")}</td><td>{p.get("field_value")}</td>'
                f'<td>{p.get("source_document")}</td><td>{p.get("page_number")}</td>'
                f'<td>{p.get("extraction_method")}</td>'
                f'<td>{round(p.get("confidence_score",0)*100)}%</td></tr>')

    def fmt_cf_row(cf):
        return (f'<tr><td><strong>{cf.get("label")}</strong></td>'
                f'<td>{cf.get("current_value")}</td><td>{cf.get("target_value")}</td>'
                f'<td>+{cf.get("score_impact")} pts</td>'
                f'<td>{cf.get("estimated_action")}</td>'
                f'<td>{cf.get("feasibility")}</td>'
                f'<td>{cf.get("implementation_timeline")}</td></tr>')

    def fmt_checklist_row(item):
        icon = "✅" if item["status"] == "PASS" else ("❌" if item["status"] == "FAIL" else "⏳")
        return (f'<tr><td>{icon} {item["requirement"]}</td>'
                f'<td>{item.get("value","")}</td><td>{item.get("notes","")}</td></tr>')

    gst_quarters_html = ""
    for q in gst_recon.get("quarters", []):
        flag_style = 'style="background:#fef2f2"' if q.get("flagged") else ""
        gst_quarters_html += (
            f'<tr {flag_style}><td>{q["quarter"]}</td>'
            f'<td>₹{q["gstr2a_itc_available"]:.2f}L</td>'
            f'<td>₹{q["gstr3b_itc_claimed"]:.2f}L</td>'
            f'<td style="color:{"#dc2626" if q["flagged"] else "#16a34a"}">'
            f'{q["variance_pct"]:+.1f}%</td>'
            f'<td>{"🚨 FLAGGED" if q["flagged"] else "✅ OK"}</td></tr>'
        )

    top_buyers_html = ""
    for b in buyer_conc.get("top_buyers", [])[:5]:
        flag_style = 'style="background:#fef2f2"' if b.get("concentration_risk_flag") else ""
        top_buyers_html += (
            f'<tr {flag_style}><td>{b.get("buyer_name") or b.get("buyer_gstin")}</td>'
            f'<td>{b.get("buyer_gstin")}</td>'
            f'<td>₹{b.get("invoice_total",0):.1f}L</td>'
            f'<td><strong>{b.get("pct_of_revenue",0):.1f}%</strong></td>'
            f'<td>{"🚨" if b.get("concentration_risk_flag") else "✅"}</td></tr>'
        )

    cf_section_html = ""
    if counterfactuals.get("counterfactuals"):
        cf_rows = "".join(fmt_cf_row(cf) for cf in counterfactuals["counterfactuals"])
        decision_style = {"APPROVE": "background:#dcfce7", "REJECT": "background:#fef2f2",
                          "CONDITIONAL_APPROVAL": "background:#fef9c3",
                          "CONDITIONAL": "background:#fef9c3"}.get(rec, "")
        cf_section_html = f"""
        <div class="section" style="{decision_style};border-radius:8px;padding:12px;margin-bottom:12px">
          <strong>Decision: {rec}</strong> |
          Score: {counterfactuals.get('current_score')}/100 |
          Threshold: {counterfactuals.get('approve_threshold')} |
          Gap: {counterfactuals.get('gap')} points
          {f'<br><em>{counterfactuals.get("buffer_message","")}</em>' if counterfactuals.get("buffer_message") else ""}
        </div>
        <table class="data-table">
          <thead><tr><th>Factor</th><th>Current</th><th>Target</th>
          <th>Score Impact</th><th>Action Required</th><th>Feasibility</th><th>Timeline</th></tr></thead>
          <tbody>{"".join(fmt_cf_row(cf) for cf in counterfactuals["counterfactuals"])}</tbody>
        </table>"""

    return f"""<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<style>
  body {{ font-family: Arial, sans-serif; font-size: 11px; color: #1f2937; margin: 20px; line-height: 1.4; }}
  h1 {{ font-size: 18px; color: #1e40af; border-bottom: 3px solid #1e40af; padding-bottom: 8px; }}
  h2 {{ font-size: 14px; color: #1e40af; border-bottom: 1px solid #dbeafe; padding-bottom: 4px; margin-top: 20px; }}
  h3 {{ font-size: 12px; color: #374151; }}
  .header-box {{ background: #eff6ff; border: 1px solid #bfdbfe; border-radius: 6px; padding: 12px; margin-bottom: 16px; }}
  .decision-badge {{ display:inline-block; padding: 6px 16px; border-radius: 4px;
                     background: {rec_color}; color: white; font-weight: bold; font-size: 14px; }}
  .section {{ margin-bottom: 16px; }}
  .data-table {{ width:100%; border-collapse:collapse; font-size:10px; margin-bottom:12px; }}
  .data-table th {{ background:#1e40af; color:white; padding:5px 8px; text-align:left; }}
  .data-table td {{ padding:4px 8px; border-bottom:1px solid #e5e7eb; }}
  .data-table tr:hover {{ background:#f9fafb; }}
  .kv-grid {{ display:grid; grid-template-columns:1fr 1fr; gap:8px; margin-bottom:8px; }}
  .kv {{ background:#f9fafb; border-radius:4px; padding:6px 10px; }}
  .kv-label {{ font-size:9px; color:#6b7280; text-transform:uppercase; }}
  .kv-value {{ font-size:13px; font-weight:bold; color:#111827; }}
  .score-bar {{ background:#e5e7eb; border-radius:4px; height:12px; margin:4px 0; }}
  .score-fill {{ height:12px; border-radius:4px; background:#1e40af; }}
  @page {{ size: A4; margin: 1.5cm; }}
</style>
</head>
<body>

<div class="header-box">
  <h1>CREDIT APPRAISAL MEMORANDUM</h1>
  <table style="width:100%"><tr>
    <td><strong>Borrower:</strong> {company.get("name","N/A")}</td>
    <td><strong>CIN:</strong> {company.get("cin","N/A")}</td>
    <td><strong>Date:</strong> {datetime.utcnow().strftime("%d %b %Y")}</td>
  </tr><tr>
    <td><strong>GSTIN:</strong> {company.get("gstin","N/A")}</td>
    <td><strong>Sector:</strong> {company.get("sector","N/A")}</td>
    <td><strong>Prepared by:</strong> IntelliCredit AI v1.0</td>
  </tr></table>
  <div style="margin-top:10px">
    <span class="decision-badge">{rec.replace("_"," ")}</span>
    &nbsp;&nbsp;
    <strong>Score: {risk_scores.get("final_score",0):.0f}/100</strong>
    &nbsp;&nbsp;
    <strong>Risk: {risk_scores.get("risk_category","N/A")}</strong>
    &nbsp;&nbsp;
    <strong>Loan: ₹{loan_terms.get("approved_amount",0):.0f}L @ {loan_terms.get("interest_rate",0):.1f}%</strong>
  </div>
</div>

<!-- SECTION 1: EXECUTIVE SUMMARY -->
<h2>1. Executive Summary</h2>
<div class="section">
  <div class="kv-grid">
    <div class="kv"><div class="kv-label">Loan Requested</div>
      <div class="kv-value">₹{data.get("loan_requested",0):.0f}L</div></div>
    <div class="kv"><div class="kv-label">Loan Approved</div>
      <div class="kv-value">₹{loan_terms.get("approved_amount",0):.0f}L</div></div>
    <div class="kv"><div class="kv-label">Interest Rate</div>
      <div class="kv-value">{loan_terms.get("interest_rate",0):.2f}% p.a.</div></div>
    <div class="kv"><div class="kv-label">Tenor</div>
      <div class="kv-value">{loan_terms.get("tenor_months",0)} months</div></div>
    <div class="kv"><div class="kv-label">Credit Score</div>
      <div class="kv-value">{risk_scores.get("final_score",0):.0f}/100</div></div>
    <div class="kv"><div class="kv-label">Default Prob (12m)</div>
      <div class="kv-value">{risk_scores.get("default_probability_12m",0):.1f}%</div></div>
  </div>
  <p><strong>Risk Flags:</strong> {len(flags)} flags detected
    ({len([f for f in flags if f.get("severity")=="CRITICAL"])} CRITICAL,
     {len([f for f in flags if f.get("severity")=="HIGH"])} HIGH)</p>
</div>

<!-- SECTION 2: COMPANY BACKGROUND -->
<h2>2. Company Background &amp; Business Overview</h2>
<div class="section">
  <p>{narratives.get("company_background","")}</p>
  <table class="data-table">
    <tr><th>Field</th><th>Value</th></tr>
    <tr><td>CIN</td><td>{company.get("cin","N/A")}</td></tr>
    <tr><td>GSTIN</td><td>{company.get("gstin","N/A")}</td></tr>
    <tr><td>PAN</td><td>{company.get("pan","N/A")}</td></tr>
    <tr><td>Sector</td><td>{company.get("sector","N/A")}</td></tr>
    <tr><td>Registered Address</td><td>{company.get("registered_address","N/A")}</td></tr>
    <tr><td>Loan Purpose</td><td>{data.get("purpose","N/A")}</td></tr>
  </table>
</div>

<!-- SECTION 3: FINANCIAL ANALYSIS -->
<h2>3. Financial Analysis</h2>
<div class="section">
  <p>{narratives.get("financial_narrative","")}</p>
  <table class="data-table">
    <thead><tr><th>Ratio</th><th>Value</th><th>Benchmark</th><th>Status</th></tr></thead>
    <tbody>
      <tr><td>DSCR</td><td>{ratios.get("dscr","N/A")}</td><td>≥ 1.25</td>
        <td>{"✅" if ratios.get("dscr") and ratios["dscr"] >= 1.25 else "❌"}</td></tr>
      <tr><td>D/E Ratio</td><td>{ratios.get("de_ratio","N/A")}</td><td>≤ 3.0</td>
        <td>{"✅" if ratios.get("de_ratio") and ratios["de_ratio"] <= 3.0 else "❌"}</td></tr>
      <tr><td>Current Ratio</td><td>{ratios.get("current_ratio","N/A")}</td><td>≥ 1.1</td>
        <td>{"✅" if ratios.get("current_ratio") and ratios["current_ratio"] >= 1.1 else "⚠️"}</td></tr>
      <tr><td>Net Profit Margin</td><td>{f'{ratios.get("net_profit_margin",0)*100:.1f}%' if ratios.get("net_profit_margin") else "N/A"}</td>
        <td>Industry var.</td><td>—</td></tr>
      <tr><td>EBITDA Margin</td><td>{f'{ratios.get("ebitda_margin",0)*100:.1f}%' if ratios.get("ebitda_margin") else "N/A"}</td>
        <td>Industry var.</td><td>—</td></tr>
      <tr><td>Interest Coverage</td><td>{ratios.get("interest_coverage","N/A")}</td>
        <td>≥ 2.0</td><td>—</td></tr>
    </tbody>
  </table>
</div>

<!-- SECTION 4: GST RECONCILIATION -->
<h2>4. GST Reconciliation — GSTR-2A vs GSTR-3B</h2>
<div class="section">
  {'<div style="background:#fef2f2;border:1px solid #fca5a5;border-radius:4px;padding:8px;margin-bottom:8px"><strong>🚨 ITC FRAUD SUSPECTED</strong> — Total suspect ITC: ₹' + str(gst_recon.get("total_suspect_itc_lakhs",0)) + 'L</div>' if gst_recon.get("itc_fraud_suspected") else '<div style="background:#dcfce7;border:1px solid #86efac;border-radius:4px;padding:8px;margin-bottom:8px">✅ GST reconciliation CLEAN — no material variance detected</div>'}
  <table class="data-table">
    <thead><tr><th>Quarter</th><th>GSTR-2A ITC Available</th>
    <th>GSTR-3B ITC Claimed</th><th>Variance %</th><th>Status</th></tr></thead>
    <tbody>{gst_quarters_html or "<tr><td colspan='5'>No GST data available</td></tr>"}</tbody>
  </table>
</div>

<!-- SECTION 5: BUYER CONCENTRATION -->
<h2>5. Buyer Revenue Concentration Analysis</h2>
<div class="section">
  {'<div style="background:#fef2f2;border:1px solid #fca5a5;border-radius:4px;padding:8px;margin-bottom:8px"><strong>🚨 SINGLE BUYER DEPENDENCY</strong> — Top buyer: ' + str(buyer_conc.get("top_buyer_pct",0)) + '% of revenue</div>' if buyer_conc.get("single_buyer_dependency") else ('<div style="background:#fef9c3;border:1px solid #fde047;border-radius:4px;padding:8px;margin-bottom:8px">⚠️ HIGH BUYER CONCENTRATION — Top 3 buyers: ' + str(buyer_conc.get("top3_concentration_pct",0)) + '%</div>' if buyer_conc.get("high_concentration") else '<div style="background:#dcfce7;border:1px solid #86efac;border-radius:4px;padding:8px;margin-bottom:8px">✅ Buyer concentration HEALTHY</div>')}
  <table class="data-table">
    <thead><tr><th>Buyer Name</th><th>GSTIN</th>
    <th>Revenue (₹L)</th><th>% of Total</th><th>Flag</th></tr></thead>
    <tbody>{top_buyers_html or "<tr><td colspan='5'>No buyer data available</td></tr>"}</tbody>
  </table>
</div>

<!-- SECTION 6: RISK ASSESSMENT -->
<h2>6. Risk Assessment — Five-Cs Framework</h2>
<div class="section">
  <p>{narratives.get("risk_narrative","")}</p>
  <table class="data-table">
    <thead><tr><th>C-Factor</th><th>Score (0-10)</th><th>Weight</th><th>Contribution</th><th>Explanation</th></tr></thead>
    <tbody>
      <tr><td>Character</td><td>{risk_scores.get("character",0):.1f}</td><td>25%</td>
        <td>{risk_scores.get("character",0)*2.5:.1f}</td>
        <td>{risk_scores.get("character_explanation","")}</td></tr>
      <tr><td>Capacity</td><td>{risk_scores.get("capacity",0):.1f}</td><td>30%</td>
        <td>{risk_scores.get("capacity",0)*3.0:.1f}</td>
        <td>{risk_scores.get("capacity_explanation","")}</td></tr>
      <tr><td>Capital</td><td>{risk_scores.get("capital",0):.1f}</td><td>20%</td>
        <td>{risk_scores.get("capital",0)*2.0:.1f}</td>
        <td>{risk_scores.get("capital_explanation","")}</td></tr>
      <tr><td>Collateral</td><td>{risk_scores.get("collateral",0):.1f}</td><td>15%</td>
        <td>{risk_scores.get("collateral",0)*1.5:.1f}</td>
        <td>{risk_scores.get("collateral_explanation","")}</td></tr>
      <tr><td>Conditions</td><td>{risk_scores.get("conditions",0):.1f}</td><td>10%</td>
        <td>{risk_scores.get("conditions",0)*1.0:.1f}</td>
        <td>{risk_scores.get("conditions_explanation","")}</td></tr>
    </tbody>
    <tfoot><tr style="background:#1e40af;color:white">
      <td><strong>TOTAL</strong></td>
      <td colspan="2"><strong>Score: {risk_scores.get("final_score",0):.1f}/100</strong></td>
      <td colspan="2"><strong>{risk_scores.get("risk_category","N/A")} RISK → {rec}</strong></td>
    </tr></tfoot>
  </table>
  <h3>Risk Flags Detected</h3>
  <table class="data-table">
    <thead><tr><th>Severity</th><th>Flag Type</th><th>Description</th></tr></thead>
    <tbody>{"".join(fmt_flag_row(f) for f in flags) or "<tr><td colspan='3'>No flags detected</td></tr>"}</tbody>
  </table>
</div>

<!-- SECTION 7: CREDIT DECISION -->
<h2>7. Credit Decision &amp; Loan Terms</h2>
<div class="section">
  <div class="kv-grid">
    <div class="kv"><div class="kv-label">Recommendation</div>
      <div class="kv-value" style="color:{rec_color}">{rec.replace("_"," ")}</div></div>
    <div class="kv"><div class="kv-label">Approved Amount</div>
      <div class="kv-value">₹{loan_terms.get("approved_amount",0):.0f}L</div></div>
    <div class="kv"><div class="kv-label">Interest Rate</div>
      <div class="kv-value">{loan_terms.get("interest_rate",0):.2f}% p.a.</div></div>
    <div class="kv"><div class="kv-label">Tenor</div>
      <div class="kv-value">{loan_terms.get("tenor_months",0)} months</div></div>
  </div>
  <h3>Covenants</h3>
  <ul>{"".join(f"<li>{c}</li>" for c in loan_terms.get("covenants",[])) or "<li>N/A</li>"}</ul>
  <h3>Monitoring Triggers</h3>
  <ul>{"".join(f"<li>{t}</li>" for t in loan_terms.get("monitoring_triggers",[])) or "<li>N/A</li>"}</ul>
  <h3>RBI Compliance Checklist</h3>
  <table class="data-table">
    <thead><tr><th>Requirement</th><th>Value</th><th>Notes</th></tr></thead>
    <tbody>{"".join(fmt_checklist_row(item) for item in checklist)}</tbody>
  </table>
</div>

<!-- SECTION 8: COUNTERFACTUALS -->
<h2>8. Path to Approval — Counterfactual Recommendations</h2>
<div class="section">
  {cf_section_html or "<p>Counterfactual analysis not available.</p>"}
</div>

<!-- APPENDIX A: CHAIN OF EVIDENCE -->
<h2>Appendix A: Chain of Evidence</h2>
<div class="section">
  <p><em>Every figure in this report is traceable to its source document, page, and extraction method.</em></p>
  <table class="data-table">
    <thead><tr><th>Field</th><th>Value</th><th>Source Document</th>
    <th>Page</th><th>Method</th><th>Confidence</th></tr></thead>
    <tbody>{"".join(fmt_prov_row(p) for p in provenance[:50]) or "<tr><td colspan='6'>No provenance data available</td></tr>"}</tbody>
  </table>
  {f'<p><em>Showing 50 of {len(provenance)} total provenance records.</em></p>' if len(provenance) > 50 else ""}
</div>

<div style="margin-top:24px;padding-top:8px;border-top:1px solid #e5e7eb;font-size:9px;color:#6b7280">
  <em>Generated by IntelliCredit AI v1.0 | IIT Hyderabad Hackathon 2025 | 
  {datetime.utcnow().strftime("%d %b %Y %H:%M")} UTC | 
  All figures sourced from uploaded documents and Sandbox.co.in API. 
  LLM used only for narrative sections. Numbers extracted deterministically.</em>
</div>
</body>
</html>"""


def export_pdf(html_content: str, output_path: str) -> bool:
    """MVP: WeasyPrint removed. Saves as .html instead — viewable in browser.
    Post-hackathon: restore WeasyPrint for true PDF output."""
    try:
        html_path = output_path.replace(".pdf", ".html")
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(html_content)
        return True
    except Exception:
        return False


def export_docx(data: dict, output_path: str) -> bool:
    """Export CAM to DOCX using python-docx."""
    try:
        from docx import Document as DocxDoc
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = DocxDoc()
        company = data.get("company", {})
        risk_scores = data.get("risk_scores", {})
        decision_data = data.get("decision", {})
        loan_terms = decision_data.get("loan_terms", {})
        counterfactuals = data.get("counterfactuals", {})

        # Title
        title = doc.add_heading("CREDIT APPRAISAL MEMORANDUM", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(
            f"Borrower: {company.get('name','N/A')} | "
            f"Date: {datetime.utcnow().strftime('%d %b %Y')} | "
            f"Generated by: IntelliCredit AI v1.0"
        )

        # Executive Summary
        doc.add_heading("1. Executive Summary", 1)
        rec = risk_scores.get("decision", "N/A")
        p = doc.add_paragraph()
        run = p.add_run(f"RECOMMENDATION: {rec}")
        run.bold = True
        doc.add_paragraph(
            f"Credit Score: {risk_scores.get('final_score',0):.0f}/100 | "
            f"Risk Category: {risk_scores.get('risk_category','N/A')} | "
            f"Approved Amount: ₹{loan_terms.get('approved_amount',0):.0f}L | "
            f"Rate: {loan_terms.get('interest_rate',0):.2f}% | "
            f"Tenor: {loan_terms.get('tenor_months',0)} months"
        )

        # Five-Cs
        doc.add_heading("2. Five-Cs Risk Scores", 1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "C-Factor"; hdr[1].text = "Score (0-10)"; hdr[2].text = "Explanation"
        for c_name, weight in [("character","25%"),("capacity","30%"),
                                ("capital","20%"),("collateral","15%"),("conditions","10%")]:
            row = table.add_row().cells
            row[0].text = f"{c_name.title()} ({weight})"
            row[1].text = str(risk_scores.get(c_name, "N/A"))
            row[2].text = risk_scores.get(f"{c_name}_explanation", "")

        # Counterfactuals
        if counterfactuals.get("counterfactuals"):
            doc.add_heading("3. Path to Approval", 1)
            doc.add_paragraph(
                f"Current Score: {counterfactuals.get('current_score')}/100 | "
                f"Gap to Approval: {counterfactuals.get('gap')} points"
            )
            cf_table = doc.add_table(rows=1, cols=4)
            cf_table.style = "Table Grid"
            hdr = cf_table.rows[0].cells
            hdr[0].text = "Factor"; hdr[1].text = "Current"; hdr[2].text = "Target"; hdr[3].text = "Action"
            for cf in counterfactuals["counterfactuals"]:
                row = cf_table.add_row().cells
                row[0].text = cf.get("label","")
                row[1].text = str(cf.get("current_value",""))
                row[2].text = str(cf.get("target_value",""))
                row[3].text = cf.get("estimated_action","")

        doc.save(output_path)
        return True
    except ImportError:
        return False
    except Exception:
        return False


# ── Main entry point ──────────────────────────────────────
async def run(app_id: str) -> dict:
    t = time.time()
    await log_agent(app_id, AGENT, "RUNNING")
    await publish_event(app_id, {
        "event_type": "AGENT_STARTED",
        "agent_name": AGENT,
        "payload": {"message": "Generating Credit Appraisal Memo..."},
        "timestamp": datetime.utcnow().isoformat(),
    })

    # Gather all data
    risk_scores = await get_session(app_id, "risk_scores") or {}
    decision_data = await get_session(app_id, "decision") or {}
    gst_recon = await get_session(app_id, "gst_reconciliation") or {}
    buyer_conc = await get_session(app_id, "buyer_concentration") or {}
    counterfactuals = await get_session(app_id, "counterfactuals") or {}
    ratios_all = await get_session(app_id, "ratios") or {}

    # Get latest ratios
    latest_ratios = {}
    if ratios_all:
        latest_key = max(ratios_all.keys(), default=None)
        if latest_key:
            latest_ratios = ratios_all.get(str(latest_key), {}) or ratios_all.get(latest_key, {})

    async with _AgentSession() as session:
        # Application + Company
        app_result = await session.execute(
            select(Application).where(Application.id == app_id)
        )
        app_obj = app_result.scalar_one_or_none()

        company_obj = None
        if app_obj:
            comp_result = await session.execute(
                select(Company).where(Company.id == app_obj.company_id)
            )
            company_obj = comp_result.scalar_one_or_none()

        # Provenance
        prov_result = await session.execute(
            select(FieldProvenance).where(FieldProvenance.application_id == app_id)
        )
        provenance = [
            {
                "field_name": p.field_name,
                "field_value": p.field_value,
                "source_document": p.source_document,
                "page_number": p.page_number,
                "extraction_method": p.extraction_method,
                "confidence_score": p.confidence_score or 0,
            }
            for p in prov_result.scalars().all()
        ]

        # Flags
        flags_result = await session.execute(
            select(RiskFlag).where(RiskFlag.application_id == app_id)
        )
        flags = [
            {"flag_type": f.flag_type, "severity": f.severity, "description": f.description}
            for f in flags_result.scalars().all()
        ]

    company_dict = {}
    if company_obj:
        company_dict = {
            "name": company_obj.name,
            "cin": company_obj.cin,
            "pan": company_obj.pan,
            "gstin": company_obj.gstin,
            "sector": company_obj.sector,
            "registered_address": company_obj.registered_address,
        }

    # Generate LLM narratives
    narratives = {
        "company_background": generate_section_narrative("company_background", {
            **company_dict,
            "purpose": app_obj.purpose if app_obj else "N/A",
        }),
        "financial_narrative": generate_section_narrative("financial_narrative", {
            "dscr": latest_ratios.get("dscr"),
            "de_ratio": latest_ratios.get("de_ratio"),
            "npm": latest_ratios.get("net_profit_margin"),
            "ebitda_margin": latest_ratios.get("ebitda_margin"),
            "current_ratio": latest_ratios.get("current_ratio"),
            "flags": ", ".join(f["flag_type"] for f in flags[:5]),
        }),
        "risk_narrative": generate_section_narrative("risk_narrative", {
            "final_score": risk_scores.get("final_score"),
            "risk_category": risk_scores.get("risk_category"),
            "decision": risk_scores.get("decision"),
            "top_drivers": str(risk_scores.get("top_drivers", [])),
        }),
    }

    # Build full data dict for template
    cam_data = {
        "company": company_dict,
        "loan_requested": app_obj.loan_amount_requested if app_obj else 0,
        "purpose": app_obj.purpose if app_obj else "N/A",
        "risk_scores": risk_scores,
        "decision": decision_data,
        "gst_reconciliation": gst_recon,
        "buyer_concentration": buyer_conc,
        "counterfactuals": counterfactuals,
        "provenance": provenance,
        "flags": flags,
        "latest_ratios": latest_ratios,
        "narratives": narratives,
    }

    # Export
    html_content = build_cam_html(cam_data)
    pdf_filename = f"CAM_{app_id[:8]}_{datetime.utcnow().strftime('%Y%m%d')}.pdf"
    docx_filename = f"CAM_{app_id[:8]}_{datetime.utcnow().strftime('%Y%m%d')}.docx"
    pdf_path = str(CAM_OUTPUT_DIR / pdf_filename)
    docx_path = str(CAM_OUTPUT_DIR / docx_filename)
    html_path = str(CAM_OUTPUT_DIR / pdf_filename.replace(".pdf", ".html"))

    # Save HTML always (fallback if WeasyPrint not installed)
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html_content)

    pdf_ok = export_pdf(html_content, pdf_path)
    docx_ok = export_docx(cam_data, docx_path)

    # Update DB
    async with _AgentSession() as session:
        cam_result = await session.execute(
            select(CAMReport).where(CAMReport.application_id == app_id)
        )
        cam = cam_result.scalar_one_or_none()
        if cam:
            cam.pdf_path = pdf_path if pdf_ok else html_path
            cam.docx_path = docx_path if docx_ok else None
            cam.generated_at = datetime.utcnow()
        else:
            cam = CAMReport(
                id=str(uuid.uuid4()),
                application_id=app_id,
                pdf_path=pdf_path if pdf_ok else html_path,
                docx_path=docx_path if docx_ok else None,
                recommendation=risk_scores.get("decision"),
                loan_amount_approved=decision_data.get("loan_terms", {}).get("approved_amount"),
                interest_rate=decision_data.get("loan_terms", {}).get("interest_rate"),
                tenor_months=decision_data.get("loan_terms", {}).get("tenor_months"),
                generated_at=datetime.utcnow(),
            )
            session.add(cam)
        await session.commit()

    result = {
        "pdf_path": pdf_path if pdf_ok else html_path,
        "docx_path": docx_path if docx_ok else None,
        "html_path": html_path,
        "pdf_generated": pdf_ok,
        "docx_generated": docx_ok,
        "sections_included": [
            "executive_summary", "company_background", "financial_analysis",
            "gst_reconciliation", "buyer_concentration", "five_cs_risk",
            "credit_decision", "counterfactuals", "chain_of_evidence_appendix",
        ],
    }

    duration_ms = int((time.time() - t) * 1000)
    await log_agent(app_id, AGENT, "COMPLETED",
                    output_summary=f"CAM generated. PDF:{pdf_ok} DOCX:{docx_ok}. "
                                   f"{len(provenance)} provenance records in appendix.",
                    duration_ms=duration_ms)
    await publish_event(app_id, {
        "event_type": "AGENT_COMPLETED",
        "agent_name": AGENT,
        "payload": result,
        "timestamp": datetime.utcnow().isoformat(),
    })
    return result



